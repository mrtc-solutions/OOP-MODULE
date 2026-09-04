"""
Week 8 DOCX assignment content — OBJECT-ORIENTED ANALYSIS: BEHAVIOURAL MODELLING.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 8,
    "title": "Object-Oriented Analysis — Behavioural Modelling",
    "filename": "Week8_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 8 module and your lecture notes before attempting each question.",
        "Draw diagrams neatly (by hand or using a UML/CASE tool) with correct notation.",
        "Show the system boundary, actor associations, include/extend relationships and swimlanes where required.",
        "Use numbered messages in communication diagrams and label every sequence-diagram message.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define behavioural modelling and explain why it is important in OOAD.", "2"),
        ("Define a use case and an actor, and give one example of each.", "2"),
        ("What is a system boundary, and where are actors and use cases placed relative to it?", "2"),
        ("Differentiate between include and extend relationships, with an example of each.", "3"),
        ("What is a use-case specification, and why is the diagram alone insufficient?", "2"),
        ("List the components of a sequence diagram and explain an activation bar.", "3"),
        ("Differentiate between a sequence diagram and a communication diagram.", "2"),
        ("What is a swimlane in an activity diagram, and why is it useful?", "2"),
        ("Explain how use cases, activity diagrams and sequence diagrams work together.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "\u201cAn ATM allows customers to withdraw cash, check balances, transfer money and change "
              "their PIN. The ATM communicates with a bank server to authorise every transaction.\u201d",
         italic=True, space_after=4)
    for s in [
        "(a)  Identify the actors (including any external system).                     [2]",
        "(b)  Identify at least five use cases.                                        [3]",
        "(c)  Draw a use-case diagram with a system boundary and actor associations.   [4]",
        "(d)  Identify one include relationship (e.g. authentication) and justify it.  [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Model the successful login of a Student using the participants Student, System and "
              "Database.", space_after=4)
    for s in [
        "(a)  Draw the sequence diagram with lifelines, activation bars and messages.  [4]",
        "(b)  Number the same interaction as a communication diagram.                 [3]",
        "(c)  Explain which diagram better shows the order of messages.               [1]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA university wants an online course registration system. A student logs in, views "
              "available courses, registers for a course (checking prerequisites and capacity), makes a "
              "payment, and receives a confirmation. A lecturer logs in and enters marks.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify the actors and use cases.                                      [3]",
        "(b)  Draw a use-case diagram with a system boundary.                         [3]",
        "(c)  Draw an activity diagram for Register Course using swimlanes.           [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (9 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — course registration", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation — ellipses for use cases, a rectangle for the system boundary, dashed lifelines, and numbered communication messages.",
        "Cite any sources you consult using the format used in the Week 8 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
