"""
Build the Week 1 assignment (Word document) — Exploits University, BICT 3202.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from common import BRAND, OUTPUT

MAROON = RGBColor(0x7B, 0x11, 0x13)
DARK   = RGBColor(0x59, 0x09, 0x0C)
SOFT   = RGBColor(0xA9, 0x44, 0x42)
GOLD   = RGBColor(0xC9, 0xA2, 0x27)
INK    = RGBColor(0x2B, 0x26, 0x26)
GREY   = RGBColor(0x6B, 0x6B, 0x6B)

H_FONT = "Poppins"
B_FONT = "Open Sans"

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def _run(p, text, size=11, bold=False, italic=False, color=INK, font=B_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return r

def _para(doc, text="", size=11, bold=False, italic=False, color=INK, font=B_FONT,
          align=None, space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        _run(p, text, size=size, bold=bold, italic=italic, color=color, font=font)
    return p

def _heading(doc, text, level=1, color=MAROON, size=None, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        size = size or 15
        _run(p, text, size=size, bold=True, color=color, font=H_FONT)
        # gold underline border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), "C9A227")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        size = size or 12.5
        _run(p, text, size=size, bold=True, color=color, font=H_FONT)
    return p

def _gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def _question(doc, num, text, marks):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _run(p, f"{num}.  ", size=11, bold=True, color=MAROON, font=H_FONT)
    _run(p, text, size=11, color=INK)
    _run(p, f"   [{marks} marks]", size=10, italic=True, color=SOFT)
    return p

def _sub(doc, text, indent=0.35):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(5)
    _run(p, text, size=11, color=INK)
    return p

def build():
    doc = Document()
    # page setup A4
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = B_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    seal = os.path.join(BRAND, "eu_seal.png")

    # ---------------- COVER ----------------
    for _ in range(2):
        _para(doc, "", space_after=0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(seal, width=Inches(1.7))
    _para(doc, "EXPLOITS UNIVERSITY", size=28, bold=True, color=MAROON, font=H_FONT,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    _para(doc, "BSc INFORMATION COMMUNICATION AND TECHNOLOGY", size=13, bold=True,
          color=SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _gold_rule(doc)
    _para(doc, "BICT 3202 · OBJECT-ORIENTED ANALYSIS AND DESIGN", size=14, bold=True,
          color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    _para(doc, "WEEK 1 ASSIGNMENT", size=30, bold=True, color=MAROON, font=H_FONT,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(doc, "Business Needs, Software Needs and Object Technology for Business",
          size=14, italic=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    # student details table
    t = doc.add_table(rows=5, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.columns[0].width = Inches(2.4)
    t.columns[1].width = Inches(3.9)
    rows = [("Student Name", ""), ("Student ID", ""), ("Programme / Year", "BSc ICT · Year 3"),
            ("Course / Week", "BICT 3202 · Week 1 of 16"), ("Date Submitted", "")]
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Inches(2.4); c1.width = Inches(3.9)
        _cell_margins(c0); _cell_margins(c1)
        _set_cell_bg(c0, "F3E4E4")
        p0 = c0.paragraphs[0]
        _run(p0, k, size=11, bold=True, color=MAROON, font=H_FONT)
        p1 = c1.paragraphs[0]
        _run(p1, v, size=11, color=INK)

    _para(doc, "", space_after=6)
    _para(doc, "Instructor: Francis Fweta — ICT Lecturer", size=12, bold=True,
          color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, "Exploits University · Week 1 of 16", size=11, italic=True, color=GREY,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------------- INSTRUCTIONS ----------------
    _heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 1 module and your lecture notes before attempting each question.",
        "Where a question asks you to “identify”, write clear, specific items — not vague single words.",
        "Support your answers with short examples. Reference ISO/IEC/IEEE 29148:2018 where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        _run(p, it, size=11, color=INK)

    # ---------------- SECTION A ----------------
    _heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define a business need.", "2"),
        ("What is a business problem, and how does it differ from a business need?", "2"),
        ("What is a business objective? Give one example.", "2"),
        ("Define a stakeholder and give two examples for a university registration system.", "2"),
        ("Distinguish between functional and non-functional requirements.", "2"),
        ("What is a business rule? Give one example.", "2"),
        ("Define object technology in your own words.", "2"),
        ("What is an object, and what are its three parts?", "2"),
        ("Differentiate between a class and an object.", "2"),
        ("List two advantages and two limitations of object-oriented approaches.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        _question(doc, i, q, m)

    # ---------------- SECTION B ----------------
    _heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    _para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    _para(doc, "A hospital's patients wait several hours before being attended to. The hospital "
               "management wants to improve the process using an information system. As an object-oriented "
               "analyst, identify:", space_after=4)
    for sub in [
        "(a)  the business problem;                                    [2]",
        "(b)  two business needs;                                      [2]",
        "(c)  two business objectives;                                 [2]",
        "(d)  five stakeholders;                                       [2]",
        "(e)  five functional requirements;                            [2]",
        "(f)  five candidate objects.                                  [2]",
    ]:
        _sub(doc, sub)
    _para(doc, "", space_after=2)
    _para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    _para(doc, "A university wants an online examination management system. Identify:", space_after=4)
    for sub in [
        "(a)  the business need;                                       [1]",
        "(b)  four stakeholders;                                       [2]",
        "(c)  three functional requirements;                           [1]",
        "(d)  three non-functional requirements;                       [2]",
        "(e)  two business rules;                                      [1]",
        "(f)  four candidate objects.                                  [1]",
    ]:
        _sub(doc, sub)

    # ---------------- SECTION C ----------------
    _heading(doc, "Section C — Case Study  (10 marks)", level=1)
    _para(doc, "“A retail organisation currently records sales using handwritten receipts. Management "
               "wants to introduce a computerised sales management system.”", italic=True, space_after=6)
    _para(doc, "As an object-oriented systems analyst, explain how you would move from understanding the "
               "business need to identifying potential objects for the proposed system. Your answer should "
               "discuss the current business problem, business objectives, business needs, stakeholders, "
               "requirements elicitation, functional and non-functional requirements, business rules, and "
               "finally the candidate objects together with their state and behaviour.", space_after=6)

    # ---------------- MARKING SCHEME ----------------
    _heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        _set_cell_bg(head[j], "7B1113")
        _cell_margins(head[j])
        p = head[j].paragraphs[0]
        _run(p, htxt, size=11, bold=True, color=RGBColor(0xFA, 0xF6, 0xEC), font=H_FONT)
    body = [
        ("A", "Short-answer questions (10 × 2)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            _cell_margins(cells[j])
            if i == 4:
                _set_cell_bg(cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            _run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    _para(doc, "", space_after=4)
    _heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Cite any sources you consult using the format used in the Week 1 module references.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        _run(p, it, size=11, color=INK)

    _gold_rule(doc)
    _para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
          size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # footer with page number
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "Exploits University · BICT 3202 OOAD — Week 1 Assignment  |  Page ", size=9, color=GREY)
    run = fp.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(9); run.font.color.rgb = GREY

    out = os.path.join(OUTPUT, "week1")
    os.makedirs(out, exist_ok=True)
    path = os.path.join(out, "Week1_Assignment_BICT3202_OOAD.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    print("Built:", build())
