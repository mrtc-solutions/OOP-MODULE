"""
Week 14 DOCX assignment content — INTEGRATING OBJECT, DYNAMIC AND BEHAVIOURAL MODELS INTO OOD.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 14,
    "title": "Integrating Object, Dynamic and Behavioural Models into Object-Oriented Design",
    "filename": "Week14_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 14 module and your lecture notes before attempting each question.",
        "Use the university course registration case study for your examples.",
        "Design classes must show visibility, data types, parameters and return types.",
        "Show how your models are consistent: messages must match operations.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define object-oriented design and differentiate it from object-oriented analysis.", "3"),
        ("Explain why analysis models must be transformed into design models.", "2"),
        ("Define a design class and list four details it adds over an analysis class.", "3"),
        ("Explain the meaning of the UML visibility symbols +, −, # and ~.", "2"),
        ("Differentiate attributes from operations, with an example.", "2"),
        ("Explain how a sequence diagram helps refine class operations.", "3"),
        ("What is model consistency, and why is it important?", "2"),
        ("Explain traceability using the \u201cRegister for Course\u201d requirement.", "3"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the requirement \u201ca student registers for an available course; the system "
              "verifies prerequisites and capacity, creates a registration and sends confirmation.\u201d",
         space_after=4)
    for s in [
        "(a)  Identify the analysis classes and the design classes.                 [3]",
        "(b)  Refine Student and Course into design classes with attributes and operations.  [3]",
        "(c)  Draw the sequence diagram, showing the controller, service and repository.  [3]",
        "(d)  Draw the state model for Course (Available, Full, Closed).            [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "The class diagram gives Course.checkCapacity(), but the sequence diagram sends "
              "Course.verifyCapacity().", space_after=4)
    for s in [
        "(a)  Identify the modelling problem.                                      [2]",
        "(b)  Explain why consistency between models matters.                       [3]",
        "(c)  Describe how the team should resolve the discrepancy.                [3]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA hospital wants an online appointment system. A patient selects a doctor, the system "
              "checks the doctor's availability and the patient's records, creates an appointment and "
              "sends a confirmation.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented designer:", space_after=4)
    for s in [
        "(a)  Identify the classes and their responsibilities.                     [3]",
        "(b)  Explain how the object, dynamic and behavioural models work together.  [4]",
        "(c)  Draw the appointment state model and explain the events.             [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — hospital appointment system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation, including visibility and message labels.",
        "Cite any sources you consult using the format used in the Week 14 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
