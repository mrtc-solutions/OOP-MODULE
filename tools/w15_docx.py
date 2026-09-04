"""
Week 15 DOCX assignment content — COMPLETE OOAD CASE STUDY, MODEL INTEGRATION AND CASE TOOL IMPLEMENTATION.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 15,
    "title": "Complete OOAD Case Study, Model Integration and CASE Tool Implementation",
    "filename": "Week15_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 15 module and your lecture notes before attempting each question.",
        "Use the university course registration case study for your examples.",
        "Show that your models are consistent: messages must match operations.",
        "Draw UML diagrams in correct notation, or with a CASE/UML tool.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define a CASE tool and state three ways it supports software development.", "3"),
        ("Differentiate functional from non-functional requirements, with one example each.", "3"),
        ("Define an actor and explain why an actor need not be a human being.", "2"),
        ("Differentiate a domain class from a technical/application class.", "3"),
        ("State the multiplicity meaning of \u201cStudent 1 — 0..* Registration\u201d.", "2"),
        ("Why should UML diagrams not be developed independently?", "3"),
        ("Explain model traceability using the \u201cRegister Course\u201d requirement.", "2"),
        ("List the six consistency checks of a model audit.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "For the university online course registration system, the requirement is: "
              "\u201ca student registers for an available course; the system verifies prerequisites and "
              "capacity, creates a registration and sends confirmation.\u201d",
         space_after=4)
    for s in [
        "(a)  Identify the actors and the functional requirements.               [2]",
        "(b)  Identify the candidate domain classes and technical classes.       [2]",
        "(c)  Draw the sequence diagram for \u201cRegister Course\u201d.                  [3]",
        "(d)  Draw the state machine for Course (Created, Available, Full, Closed).  [2]",
        "(e)  Explain how the sequence messages map to class operations.          [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "The activity diagram contains the step \u201cCheck Payment\u201d, but the class diagram has "
              "no payment-related class, operation or component.", space_after=4)
    for s in [
        "(a)  Identify the modelling problem.                                      [2]",
        "(b)  Explain how the models can be made consistent.                       [3]",
        "(c)  Describe how a CASE tool would help manage this change.              [3]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA college wants an Online Library Management System. Students can search, borrow, "
              "return and view borrowed books. Librarians can add books, remove books, register members "
              "and view overdue books.\u201d",
         italic=True, space_after=6)
    for s in [
        "(a)  Identify the actors, use cases and candidate classes.               [3]",
        "(b)  Draw the sequence diagram for \u201cBorrow Book\u201d.                       [3]",
        "(c)  Draw the state diagram for Book Copy (Available, Borrowed, Returned).  [2]",
        "(d)  Explain how the five diagrams relate to one another.                [2]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — library management system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation, including multiplicity and message labels.",
        "Cite any sources you consult using the format used in the Week 15 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
