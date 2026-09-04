"""
Week 13 DOCX assignment content — MULTILAYER SYSTEMS, REUSABLE COMPONENTS AND ARCHITECTURAL DESIGN.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 13,
    "title": "Multilayer Systems, Reusable Components and Architectural Design",
    "filename": "Week13_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 13 module and your lecture notes before attempting each question.",
        "Use the university course registration case study for your examples.",
        "Component, package and deployment diagrams must use correct UML notation.",
        "Where you critique an architecture, propose a concrete improvement.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define software architecture and distinguish it from class design.", "3"),
        ("Define a software layer and list the layers of a common multilayer architecture.", "3"),
        ("Explain the responsibilities of the Presentation and Domain layers.", "2"),
        ("Explain separation of concerns.", "2"),
        ("Define a software component and differentiate it from a class.", "3"),
        ("Explain why components need interfaces, using a payment example.", "2"),
        ("Differentiate a reusable component from an adaptable component.", "2"),
        ("State two benefits and two limitations of multilayer architecture.", "3"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Exploits University's registration system has the following classes: Student, Course, "
              "Registration, RegistrationService, StudentRepository, CourseRepository, "
              "RegistrationRepository, and web/mobile interfaces.", space_after=4)
    for s in [
        "(a)  Organise these into a layered architecture, naming each layer.          [4]",
        "(b)  Draw a UML component diagram for the system.                            [4]",
        "(c)  Identify two components that could be reused by other applications, and explain why.  [2]",
        "(d)  Explain how the design keeps coupling low.                              [2]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "A StudentController contains HTML generation, SQL queries, password validation, course "
              "registration, email sending, fee calculation and PDF generation.", space_after=4)
    for s in [
        "(a)  Explain what is wrong with this design.                                [2]",
        "(b)  State whether cohesion is high or low, and why.                        [2]",
        "(c)  Propose a better architecture, naming the layers and services.        [4]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA bank wants an online banking system with a customer interface, authentication, "
              "account management, transactions, payments and notifications. The bank expects to add new "
              "channels (mobile, web, ATM) and new payment methods over time.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented architect:", space_after=4)
    for s in [
        "(a)  Identify the major components and their responsibilities.              [3]",
        "(b)  Design the interfaces and dependencies between the components.        [4]",
        "(c)  Explain how the architecture promotes reuse and adaptability.         [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — online banking system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation for components, interfaces, packages and dependencies.",
        "Cite any sources you consult using the format used in the Week 13 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
