"""
Week 4 DOCX assignment content — UML BUILDING BLOCKS, RULES, MECHANISMS AND ARCHITECTURE.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 4,
    "title": "UML Building Blocks, Rules, Mechanisms and Architecture",
    "filename": "Week4_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 4 module and your lecture notes before attempting each question.",
        "Draw diagrams neatly (by hand or using a UML/diagramming tool) with correct notation.",
        "Use correct UML notation — a diamond, triangle or dashed arrow has a specific meaning.",
        "Cite OMG UML 2.5.1 where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("State the three conceptual building blocks of UML and briefly explain each.", "3"),
        ("Name and give one example each of the four kinds of UML “things”.", "3"),
        ("List the four main UML relationship types and draw the notation for each.", "3"),
        ("Differentiate between syntax and semantics in UML.", "2"),
        ("Explain what it means for a UML model to be well-formed.", "2"),
        ("Explain the four common mechanisms of UML.", "2"),
        ("Define a stereotype, a tagged value and a constraint, with an example of each.", "3"),
        ("Explain the four UML layers M3, M2, M1 and M0.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Study the following statements about a university student registration system and, for each, "
              "identify the most appropriate UML relationship, then draw its notation:",
         space_after=4)
    for s in [
        "(a)  A Student registers for a Course.                              [2]",
        "(b)  A Student is a specialised type of User.                       [2]",
        "(c)  A Registration uses information from the Course class.         [2]",
        "(d)  MobilePayment implements the PaymentService interface.         [2]",
        "(e)  A Department contains many Lecturers (whole–part, shared).     [2]",
        "(f)  A Lecturer teaches a Course.                                   [2]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider a UML class diagram fragment for a library system with the classes Book, Member and "
              "Loan.", space_after=4)
    for s in [
        "(a)  Draw the three classes with two attributes each.              [2]",
        "(b)  Add the relationships Member—Loan and Loan—Book.              [2]",
        "(c)  Add multiplicities (Member 1 ——— 0..* Loan; Book 1 ——— 0..* Loan).  [2]",
        "(d)  Add a constraint or note stating that a member may borrow at most five books.  [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "“A university wants an online e-learning platform. Students can log in, view courses, "
              "download notes, submit assignments and view grades. Lecturers can upload notes, create "
              "assignments, mark assignments and publish grades. A submission moves from Draft to "
              "Submitted, then Marked, then Returned.”", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify the candidate classes and name two attributes for each.  [3]",
        "(b)  Describe the key relationships between the classes, naming each relationship type.  [3]",
        "(c)  Use the four-layer architecture (M0–M3) to explain where “Student”, “Submission”, “Class” and “MOF” belong.  [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation — marks are awarded for notation as well as content.",
        "Cite any sources you consult using the format used in the Week 4 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
