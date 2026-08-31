"""
Week 12 DOCX assignment content — INTEGRATED OBJECT-ORIENTED MODELLING AND DESIGN.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 12,
    "title": "Integrated Object-Oriented Modelling and Design",
    "filename": "Week12_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 12 module and your lecture notes before attempting each question.",
        "Where you draw UML diagrams, use correct notation — actors, multiplicities, states, messages.",
        "Show how your models are consistent: every class in a sequence diagram must exist in the class model.",
        "Cite the OMG UML 2.5.1 specification where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define integrated object-oriented modelling.", "2"),
        ("Explain why the different UML models of a system must be consistent.", "3"),
        ("State the main question answered by each of the use-case, class, sequence, activity and state models.", "3"),
        ("Explain traceability and why it is valuable.", "2"),
        ("Differentiate a model from a diagram, using an analogy.", "2"),
        ("Explain the purpose of packages in a large UML model.", "2"),
        ("Differentiate a sequence diagram from an activity diagram.", "3"),
        ("What is design validation?", "3"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the requirement \u201ca student can register for an available course online\u201d.",
         space_after=4)
    for s in [
        "(a)  Draw the use-case diagram.                                              [2]",
        "(b)  Draw the class diagram, including multiplicities.                       [3]",
        "(c)  Draw the sequence diagram for course registration.                      [3]",
        "(d)  Draw the state diagram for the Registration lifecycle.                  [2]",
        "(e)  Explain how the models are consistent with one another.                 [2]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "A sequence diagram contains the operation checkAvailability(), but the corresponding "
              "class diagram contains no such operation.", space_after=4)
    for s in [
        "(a)  Identify the modelling problem.                                         [2]",
        "(b)  Explain why consistency between models is important.                    [3]",
        "(c)  Describe how the designer should resolve the problem.                   [3]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA hospital wants an online appointment system. A patient can select a doctor, choose an "
              "available date and time, and book an appointment.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented analyst/designer, produce an integrated model:", space_after=4)
    for s in [
        "(a)  Identify the actors and use cases.                                      [2]",
        "(b)  Identify the classes, attributes, operations and relationships.         [3]",
        "(c)  Describe the sequence of interactions and the appointment states.       [3]",
        "(d)  Explain how your models are consistent.                                  [2]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — hospital appointment system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation, including multiplicities and message labels.",
        "Cite any sources you consult using the format used in the Week 12 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
