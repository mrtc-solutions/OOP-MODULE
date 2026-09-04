"""
DOCX engine — brand-styled assignment builder, parameterized by week.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from common import BRAND, OUTPUT

MAROON = RGBColor(0x7B, 0x11, 0x13)
DARK   = RGBColor(0x59, 0x09, 0x0C)
SOFT   = RGBColor(0xA9, 0x44, 0x42)
GOLD   = RGBColor(0xC9, 0xA2, 0x27)
INK    = RGBColor(0x2B, 0x26, 0x26)
GREY   = RGBColor(0x6B, 0x6B, 0x6B)
CREAM  = RGBColor(0xFA, 0xF6, 0xEC)

H_FONT = "Poppins"
B_FONT = "Open Sans"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def run(p, text, size=11, bold=False, italic=False, color=INK, font=B_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return r


def para(doc, text="", size=11, bold=False, italic=False, color=INK, font=B_FONT,
         align=None, space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, size=size, bold=bold, italic=italic, color=color, font=font)
    return p


def heading(doc, text, level=1, color=MAROON, size=None, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        size = size or 15
        run(p, text, size=size, bold=True, color=color, font=H_FONT)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), "C9A227")
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        size = size or 12.5
        run(p, text, size=size, bold=True, color=color, font=H_FONT)
    return p


def gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A227")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def question(doc, num, text, marks):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, f"{num}.  ", size=11, bold=True, color=MAROON, font=H_FONT)
    run(p, text, size=11, color=INK)
    run(p, f"   [{marks} marks]", size=10, italic=True, color=SOFT)
    return p


def sub(doc, text, indent=0.35):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(5)
    run(p, text, size=11, color=INK)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run(p, text, size=11, color=INK)
    return p


def cover(doc, cfg):
    for _ in range(2):
        para(doc, "", space_after=0)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(BRAND, "eu_seal.png"), width=Inches(1.7))
    para(doc, "EXPLOITS UNIVERSITY", size=28, bold=True, color=MAROON, font=H_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    para(doc, "BSc INFORMATION COMMUNICATION AND TECHNOLOGY", size=13, bold=True,
         color=SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    gold_rule(doc)
    para(doc, "BICT 3202 · OBJECT-ORIENTED ANALYSIS AND DESIGN", size=14, bold=True,
         color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, f"WEEK {cfg['week']} ASSIGNMENT", size=30, bold=True, color=MAROON, font=H_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, cfg["title"], size=14, italic=True, color=INK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    t = doc.add_table(rows=5, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    rows = [("Student Name", ""), ("Student ID", ""), ("Programme / Year", "BSc ICT · Year 3"),
            ("Course / Week", f"BICT 3202 · Week {cfg['week']} of 16"), ("Date Submitted", "")]
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Inches(2.4); c1.width = Inches(3.9)
        cell_margins(c0); cell_margins(c1)
        set_cell_bg(c0, "F3E4E4")
        p0 = c0.paragraphs[0]
        run(p0, k, size=11, bold=True, color=MAROON, font=H_FONT)
        p1 = c1.paragraphs[0]
        run(p1, v, size=11, color=INK)

    para(doc, "", space_after=6)
    para(doc, "Instructor: Francis Fweta — ICT Lecturer", size=12, bold=True,
         color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, f"Exploits University · Week {cfg['week']} of 16", size=11, italic=True,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_footer(doc):
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(fp, "Exploits University · BICT 3202 OOAD — Assignment  |  Page ", size=9, color=GREY)
    r = fp.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)
    r.font.size = Pt(9); r.font.color.rgb = GREY


def build(cfg, content_builder):
    doc = Document()
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    normal = doc.styles["Normal"]
    normal.font.name = B_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    H = {
        "para": para, "heading": heading, "gold_rule": gold_rule, "question": question,
        "sub": sub, "bullet": bullet, "set_cell_bg": set_cell_bg,
        "cell_margins": cell_margins, "run": run,
        "MAROON": MAROON, "DARK": DARK, "SOFT": SOFT, "INK": INK, "GREY": GREY, "CREAM": CREAM,
        "H_FONT": H_FONT, "B_FONT": B_FONT,
    }

    cover(doc, cfg)
    content_builder(doc, H)
    add_footer(doc)

    out = os.path.join(OUTPUT, f"week{cfg['week']}")
    os.makedirs(out, exist_ok=True)
    path = os.path.join(out, cfg["filename"])
    doc.save(path)
    return path
