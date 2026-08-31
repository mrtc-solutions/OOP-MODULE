"""
Week 11 DOCX assignment content — REFINING THE OBJECT-ORIENTED DESIGN.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 11,
    "title": "Refining the Object-Oriented Design",
    "filename": "Week11_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 11 module and your lecture notes before attempting each question.",
        "Use the university course registration case study for your examples.",
        "UML notation must be correct — visibility (−, +), types, parameters and return values.",
        "Where you critique a design, propose a concrete improvement, not just a description.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define design refinement and explain its purpose.", "2"),
        ("Differentiate an analysis class from a design class, with an example.", "3"),
        ("Define entity, boundary and control classes, giving one example of each.", "3"),
        ("Explain why operation signatures are important, using registerCourse(course : Course) : Registration.", "2"),
        ("Differentiate cohesion from coupling.", "2"),
        ("Differentiate an abstract class from an interface.", "3"),
        ("Explain the HAS / IS-A test for choosing between association and inheritance.", "2"),
        ("What is dependency injection, and how does it improve a design?", "3"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the requirement \u201cstudents can register for courses online\u201d.", space_after=4)
    for s in [
        "(a)  Refine Student, Course and Registration from analysis classes into design classes (attributes and operations).  [4]",
        "(b)  Assign responsibilities: which class should check course availability, and why?   [2]",
        "(c)  Design a RegistrationService and explain why it is needed.                        [3]",
        "(d)  Show how an interface could reduce coupling between the service and the data layer.  [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "The following class is proposed:", space_after=4)
    para(doc, "StudentSystem: login(), register(), payFees(), sendEmail(), sendSMS(), generateResults(), connectDatabase()",
         size=10.5, font="Consolas", space_after=4)
    for s in [
        "(a)  Explain whether the class is highly cohesive, and why.                        [2]",
        "(b)  Identify at least five classes/services to replace it.                       [3]",
        "(c)  Explain how your redesign improves cohesion and coupling.                    [3]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA university wants an online payment system supporting bank payments, mobile-money "
              "payments and card payments. New payment methods should be added in the future with "
              "minimal changes to existing code.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented designer:", space_after=4)
    for s in [
        "(a)  Produce a class/interface model for the system.                            [4]",
        "(b)  Explain the responsibilities of each element and how coupling is reduced. [3]",
        "(c)  Explain how the design supports reuse and change.                         [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — online payment system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation, including visibility, types and multiplicities.",
        "Cite any sources you consult using the format used in the Week 11 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
