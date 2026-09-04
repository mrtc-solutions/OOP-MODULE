"""
Week 2 DOCX assignment content — PRINCIPLES OF OBJECT MODELLING.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 2,
    "title": "Principles of Object Modelling",
    "filename": "Week2_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    # instructions
    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 2 module and your lecture notes before attempting each question.",
        "Where a question asks you to “identify”, write clear, specific items and justify each choice.",
        "Cite ISO/IEC/IEEE 29148:2018 and OMG UML 2.5.1 where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define object modelling.", "2"),
        ("Define an object model.", "2"),
        ("Why are object models useful in systems analysis?", "2"),
        ("Explain the relationship between an object, a class and an object model.", "2"),
        ("Define abstraction and give one example from object modelling.", "2"),
        ("Define encapsulation and explain why it matters for a bank account.", "2"),
        ("What is object identity? How does it differ from an attribute?", "2"),
        ("Distinguish between object state and object behaviour.", "2"),
        ("What is object identification? Why should not every noun become a class?", "2"),
        ("Name the standard language used to express object models, and state the current published specification version.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "“A customer uses a mobile phone to purchase a product from an online shop. The customer "
              "adds the product to a shopping cart and pays using mobile money. The system generates an "
              "invoice.”", italic=True, space_after=4)
    para(doc, "Identify:", space_after=4)
    for s in [
        "(a)  six candidate objects;                                   [3]",
        "(b)  two attributes for three of the objects;                 [3]",
        "(c)  two behaviours for two of the objects;                   [2]",
        "(d)  three relationships;                                     [2]",
        "(e)  two external systems, and state whether “Mobile Phone”, “Mobile Money” and “Invoice” should be objects — with reasons.   [2]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "A hospital wants an electronic appointment management system. Patients request "
              "appointments with doctors; doctors have specialties; receptionists manage appointments.", space_after=4)
    for s in [
        "(a)  five candidate objects;                                  [2]",
        "(b)  attributes for Patient and Appointment;                  [2]",
        "(c)  three behaviours;                                        [2]",
        "(d)  two relationships, and one piece of information to be excluded through abstraction.   [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "“A university wants to develop an online course registration system. Students select "
              "courses and submit a registration; lecturers approve registrations; the system enforces "
              "prerequisites and records payments.”", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify ten candidate objects.                          [2]",
        "(b)  Give five attributes for appropriate objects and five behaviours.   [3]",
        "(c)  Give five relationships between objects.                 [2]",
        "(d)  Explain how abstraction, encapsulation and modularity would help you keep this model simple and maintainable.   [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (10 × 2)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Cite any sources you consult using the format used in the Week 2 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
