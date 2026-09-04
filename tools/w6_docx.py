"""
Week 6 DOCX assignment content — OBJECT-ORIENTED ANALYSIS: OBJECT MODELLING II.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 6,
    "title": "Object-Oriented Analysis — Object Modelling II",
    "filename": "Week6_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 6 module and your lecture notes before attempting each question.",
        "Draw diagrams neatly (by hand or using a UML/CASE tool) with correct UML notation.",
        "Use a hollow diamond for aggregation and a filled diamond for composition.",
        "Justify every modelling decision — marks are awarded for reasoning, not only notation.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define aggregation and explain the whole-part relationship.", "2"),
        ("Differentiate between ordinary association and aggregation, with an example of each.", "3"),
        ("Differentiate between aggregation and composition, with an example of each.", "3"),
        ("What is inheritance? How does it relate to UML generalization?", "2"),
        ("What features does a subclass inherit from its superclass? Give an example.", "2"),
        ("Differentiate between inheritance (is-a) and aggregation (has-a).", "2"),
        ("Define polymorphism and give one real-world example.", "2"),
        ("Explain method overriding and how it relates to polymorphism.", "2"),
        ("Why does polymorphism support extensibility in a software system?", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "\u201cA university has departments and lecturers. Each department has many lecturers, and a "
              "lecturer belongs to one department. A lecturer can move to another department. A course is "
              "made up of course materials such as slides, notes and assessments, which belong to that "
              "course alone.\u201d", italic=True, space_after=4)
    for s in [
        "(a)  Model Department and Lecturer, and explain whether aggregation or association is appropriate.  [3]",
        "(b)  Model Course and CourseMaterial, and explain whether aggregation or composition is appropriate.  [3]",
        "(c)  Add multiplicities to both relationships.                                    [2]",
        "(d)  Draw both relationships using correct UML notation (diamonds).               [4]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the hierarchy Animal -> Dog, Cat and Bird, where each animal makes a sound.",
         space_after=4)
    for s in [
        "(a)  Draw the class hierarchy using generalization.                              [2]",
        "(b)  Explain how polymorphism applies to the operation makeSound().              [2]",
        "(c)  Explain how method overriding is used to give each animal its own sound.    [2]",
        "(d)  Explain what happens if a new class, Lion, is added to the hierarchy.       [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cAn online store has customers, orders, order items and products. Payments can be made "
              "by card, mobile money or bank transfer. Each order contains one or more order items. The "
              "system is organised into customer, order, product and payment management areas.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify the candidate classes, including a Payment hierarchy.             [3]",
        "(b)  Model Order and OrderItem and justify aggregation or composition.          [3]",
        "(c)  Explain how polymorphism applies to processPayment() and why it supports extensibility.  [2]",
        "(d)  Propose a package structure (grouping) for the system.                     [2]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (9 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — online store", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation — hollow diamond for aggregation, filled diamond for composition.",
        "Cite any sources you consult using the format used in the Week 6 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
