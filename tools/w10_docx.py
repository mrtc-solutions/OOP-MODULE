"""
Week 10 DOCX assignment content — OBJECT-ORIENTED DESIGN.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 10,
    "title": "Object-Oriented Design",
    "filename": "Week10_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 10 module and your lecture notes before attempting each question.",
        "Use the university student registration case study for your examples.",
        "Diagrams (class, sequence, state) must be neat, complete and correctly labelled.",
        "Cite the OMG UML 2.5.1 specification and IBM design/analysis model guidance where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define software design and object-oriented design.", "3"),
        ("Differentiate object-oriented analysis from object-oriented design, with an example.", "3"),
        ("State four goals of good object-oriented design.", "2"),
        ("Define cohesion and coupling.", "2"),
        ("Explain why good design generally seeks high cohesion and low coupling.", "2"),
        ("Differentiate abstraction from encapsulation using a BankAccount example.", "3"),
        ("Explain the Object, Dynamic and Behavioural models.", "3"),
        ("What is traceability, and why is it valuable?", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Exploits University wants an Online Student Registration System. Students will log in, "
              "search for courses, register for courses, drop courses and view registered courses.",
         space_after=4)
    for s in [
        "(a)  Identify five possible classes for the system.                              [2]",
        "(b)  Assign at least two responsibilities to each class.                         [3]",
        "(c)  Draw a class diagram showing attributes, operations, associations and multiplicities.  [3]",
        "(d)  Explain how the Object, Dynamic and Behavioural models support one another. [4]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the requirement \u201ca student registers for a course\u201d.", space_after=4)
    for s in [
        "(a)  Draw a sequence diagram for the course registration flow.                  [4]",
        "(b)  Draw a state diagram for the Registration object.                          [2]",
        "(c)  Identify one cohesion and one coupling problem that could arise, and suggest a fix.  [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA bank wants a mobile banking application allowing customers to check balances, "
              "transfer money, pay bills and buy airtime. The system must be secure, support many "
              "simultaneous users, and allow new payment methods to be added in future.\u201d",
         italic=True, space_after=6)
    para(doc, "As an object-oriented designer:", space_after=4)
    for s in [
        "(a)  Identify five classes and their responsibilities.                          [3]",
        "(b)  Explain how the three models (object, dynamic, behavioural) would be combined.  [3]",
        "(c)  Explain how an interface and a design pattern could support adding new payment methods.  [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — mobile banking system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Hand-drawn diagrams must use a ruler and correct UML notation; software-drawn diagrams are welcome.",
        "Where you describe cohesion, coupling, abstraction or encapsulation, illustrate with the given scenario rather than definitions only.",
        "Cite any sources you consult using the format used in the Week 10 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
