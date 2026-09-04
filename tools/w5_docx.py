"""
Week 5 DOCX assignment content — OBJECT-ORIENTED ANALYSIS: OBJECT MODELLING I.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 5,
    "title": "Object-Oriented Analysis — Object Modelling I",
    "filename": "Week5_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 5 module and your lecture notes before attempting each question.",
        "Draw diagrams neatly (by hand or using a UML/CASE tool) with correct UML notation.",
        "For every association, show a meaningful name and multiplicities where required.",
        "Cite OMG UML 2.5.1 where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define Object-Oriented Analysis.", "2"),
        ("Differentiate between a class and an object, giving one example of each.", "3"),
        ("What is an attribute? Give two attributes of the class Course.", "2"),
        ("What is an operation? Give two operations of the class Student.", "2"),
        ("Differentiate between a link and an association.", "2"),
        ("What is multiplicity? Explain the meaning of 0..1, 1, 0..* and 1..*.", "3"),
        ("Define generalization and specialization.", "2"),
        ("Explain the \u201cis-a\u201d test and use it to decide whether Student \u2192 Course is a generalization.", "2"),
        ("What is a CASE tool, and why can it not replace an analyst?", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (20 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "\u201cA university has students and lecturers. Each student has a student ID, name, programme "
              "and year of study. Each lecturer has a staff ID, name and department. Students register for "
              "courses. Each course has a course code, title and credit hours. Lecturers teach courses. A "
              "student may register for many courses, while each course may have many students.\u201d",
         italic=True, space_after=4)
    for s in [
        "(a)  Identify the candidate classes.                                        [3]",
        "(b)  Identify at least two attributes for each class.                        [4]",
        "(c)  Identify one or two operations that could reasonably belong to each class.  [3]",
        "(d)  Draw a UML class diagram showing the associations, with meaningful names.  [5]",
        "(e)  Add appropriate multiplicities to the associations.                     [3]",
        "(f)  State whether generalization is appropriate, and justify your answer.   [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA university wants an Accommodation Management System to manage students, rooms, "
              "hostels, accommodation applications, payments and wardens. Students apply for rooms. A "
              "hostel contains several rooms. A student can occupy one room at a time. A warden manages a "
              "hostel. Students make accommodation payments.\u201d", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify the candidate classes (at least six).                        [3]",
        "(b)  Provide at least two attributes for each major class.                 [3]",
        "(c)  Draw the associations between the classes and add multiplicities.     [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (9 questions)", "20"),
        ("B", "Application — class diagram (parts a\u2013f)", "20"),
        ("C", "Case study — accommodation system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation — marks are awarded for notation, association names and multiplicities.",
        "Cite any sources you consult using the format used in the Week 5 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
