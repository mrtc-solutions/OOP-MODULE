"""
Week 9 DOCX assignment content — THE UNIFIED SOFTWARE DEVELOPMENT PROCESS (USDP / UP).
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 9,
    "title": "The Unified Software Development Process (USDP / Unified Process)",
    "filename": "Week9_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 9 module and your lecture notes before attempting each question.",
        "Use clear examples from the university registration case study where possible.",
        "Diagrams (phases, iterations, workflows) must be neat and correctly labelled.",
        "Cite the OMG UML 2.5.1 specification and RUP references where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define the Unified Software Development Process.", "2"),
        ("Differentiate between the Unified Process and the Rational Unified Process.", "2"),
        ("Explain the difference between iterative and incremental development, with an example of each.", "3"),
        ("State and explain the four major characteristics of the Unified Process.", "4"),
        ("List the four UP phases and state the main question answered by each.", "3"),
        ("Differentiate between a phase, an iteration, a workflow and an artefact.", "2"),
        ("What is software reuse? Give two guidelines for evaluating a reusable component.", "2"),
        ("Differentiate between generalisation and inheritance.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Exploits University wants an Online Student Management System. Students will register, "
              "view courses, pay fees and view results; lecturers will view classes, upload materials and "
              "enter marks; administrators will manage students and courses and generate reports.",
         space_after=4)
    for s in [
        "(a)  Explain what would happen during Inception, listing its major activities and outputs.  [3]",
        "(b)  Explain why Elaboration is particularly important, and give two architectural risks.   [3]",
        "(c)  Describe three iterations that could occur during Construction.                      [3]",
        "(d)  Describe the activities of Transition.                                               [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider the requirement \u201cstudents must register for courses\u201d.", space_after=4)
    for s in [
        "(a)  Explain how UP is use-case driven, using this requirement.                          [3]",
        "(b)  Explain how the Week 8 models (use case, activity, sequence) become artefacts in UP.  [3]",
        "(c)  Identify two reusable assets that could support this system and evaluate each.       [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA hospital intends to develop an online patient management system to register patients, "
              "book appointments, process payments, store medical records and generate reports. The "
              "hospital is particularly concerned about security, performance and integration with "
              "existing systems.\u201d", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Explain how UP could be used to develop the system.                               [3]",
        "(b)  Explain what should happen during Inception and why Elaboration is particularly important.  [3]",
        "(c)  Give three architectural risks and suggest three artefacts for the project.       [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — hospital system", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Where you describe UP phases, iterations and artefacts, relate them to the given scenario rather than reciting definitions only.",
        "Cite any sources you consult using the format used in the Week 9 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
