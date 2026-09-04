"""
Week 3 DOCX assignment content — MODELLING LANGUAGE: UML.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 3,
    "title": "Modelling Language — UML",
    "filename": "Week3_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 3 module and your lecture notes before attempting each question.",
        "Draw diagrams neatly (by hand or using a UML/diagramming tool) with correct notation.",
        "Cite OMG UML 2.5.1 and MOF where relevant.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define UML and state what each part of the name means.", "3"),
        ("Explain the difference between a modelling language and a development methodology, using UML and UP as examples.", "2"),
        ("Name the three conceptual building blocks of UML.", "2"),
        ("List and give one example each of the four kinds of UML “things”.", "2"),
        ("Name the four main UML relationship types.", "2"),
        ("Differentiate between syntax and semantics in UML.", "2"),
        ("Explain the four common mechanisms of UML.", "2"),
        ("Define a stereotype, a tagged value and a constraint, with an example of each.", "3"),
        ("Explain the four UML layers M3, M2, M1 and M0.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "A university wants a mobile student registration application. Students can log in, view "
              "courses, register courses, drop courses and view their registration status. The application "
              "communicates with a university server and a database.", space_after=4)
    para(doc, "For each requirement below, state the most appropriate UML diagram and justify your choice:",
         space_after=4)
    for s in [
        "(a)  Show what functions students can perform.             [2]",
        "(b)  Show the Student and Course classes and their relationship.   [2]",
        "(c)  Show the login communication between the app, server and database.   [2]",
        "(d)  Show the steps in the course-registration workflow.   [2]",
        "(e)  Show the states of a registration (Draft → Submitted → Approved).   [2]",
        "(f)  Show where the app, server and database are deployed.   [2]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Draw a simple UML class diagram for the university registration system containing Student, "
              "Course, Lecturer and Registration.", space_after=4)
    for s in [
        "(a)  Show the four classes with two attributes each.        [2]",
        "(b)  Add the relationships Student—Registration, Course—Registration and Lecturer—Course.   [2]",
        "(c)  Add sensible multiplicities to two of the relationships.   [2]",
        "(d)  Add one constraint or note expressing a business rule (e.g. prerequisites).   [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "“An online banking system allows customers to view balances, transfer money between "
              "accounts and pay bills. The system authenticates customers, communicates with a core banking "
              "server and a database, and must encrypt all transfers.”", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Explain how UML would help the analyst, developer, database designer and bank manager communicate about this system.   [3]",
        "(b)  Identify which UML diagrams you would use to show: the functions a customer can perform; the Account and Transaction classes; the order of a transfer; the states of a transaction.   [4]",
        "(c)  Use the four-layer architecture (M0–M3) to explain where “Customer”, “Account”, “Class” and “MOF” belong.   [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (9 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Cite any sources you consult using the format used in the Week 3 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
