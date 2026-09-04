"""
Week 7 DOCX assignment content — OBJECT-ORIENTED ANALYSIS: DYNAMIC MODELLING.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 7,
    "title": "Object-Oriented Analysis — Dynamic Modelling",
    "filename": "Week7_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 7 module and your lecture notes before attempting each question.",
        "Draw state-machine diagrams neatly (by hand or using a UML/CASE tool) with correct notation.",
        "Use a filled circle for the initial state and a bull's-eye for the final state; label every transition.",
        "Justify guards and every transition you include.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define dynamic modelling and explain why it is needed in an object-oriented system.", "2"),
        ("Define an event and give two examples from different categories.", "2"),
        ("Define a state and differentiate it from an event.", "3"),
        ("What is a transition? What is a guard condition? Give an example of each.", "3"),
        ("Explain entry, exit and do behaviour, with one example of each.", "3"),
        ("Define an operation and differentiate it from an event.", "2"),
        ("What is concurrency? Differentiate sequential from concurrent behaviour.", "3"),
        ("Explain why a class diagram alone is insufficient for modelling a complex Order object.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "\u201cA bank account can be Active, Blocked or Closed. The events that occur are login, "
              "wrongPIN, blockAccount, unblockAccount and closeAccount.\u201d", italic=True, space_after=4)
    for s in [
        "(a)  Identify the states and the events.                                       [2]",
        "(b)  Draw the state-machine diagram with an initial state.                      [4]",
        "(c)  Add a guard that blocks the account after three wrong PIN entries.         [2]",
        "(d)  Identify any transition that should NOT be allowed, and justify your answer.  [4]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "Consider an online payment with states Pending, Processing, Paid and Failed, and events "
              "submitPayment, paymentSuccessful and paymentFailed.", space_after=4)
    for s in [
        "(a)  Draw the state-machine diagram.                                            [3]",
        "(b)  Add guards to distinguish the successful and failed outcomes.              [2]",
        "(c)  Add entry behaviour to the Paid state (e.g. send a confirmation email).     [1]",
        "(d)  Explain, in words, what the diagram says as a story.                        [2]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA food delivery application performs the following after an order is placed: payment "
              "processing, restaurant order preparation, driver assignment and customer notification. An "
              "order moves from Placed to Preparing, then to Out for Delivery, then to Delivered (or "
              "Cancelled).\u201d", italic=True, space_after=6)
    para(doc, "As an object-oriented analyst:", space_after=4)
    for s in [
        "(a)  Identify the states and events of the Order.                              [3]",
        "(b)  Draw the Order state-machine diagram.                                     [3]",
        "(c)  Explain which activities could occur concurrently and why.                [4]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (8 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — food delivery", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "State-machine diagrams must use correct UML notation — initial state, states, labelled transitions, guards, final state.",
        "Cite any sources you consult using the format used in the Week 7 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
