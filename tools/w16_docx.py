"""
Week 16 DOCX assignment content — COMPREHENSIVE REVISION, INTEGRATED CASE STUDY AND EXAMINATION PREPARATION.
"""
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import build, para, heading, question, sub, bullet, gold_rule, run

CFG = {
    "week": 16,
    "title": "Comprehensive Revision, Integrated Case Study and Examination Preparation",
    "filename": "Week16_Assignment_BICT3202_OOAD.docx",
}


def content(doc, H):
    MAROON = H["MAROON"]; DARK = H["DARK"]; INK = H["INK"]; CREAM = H["CREAM"]
    H_FONT = H["H_FONT"]

    heading(doc, "Assignment Instructions", level=1)
    for it in [
        "Answer ALL questions in Sections A, B and C.",
        "Read the Week 16 module and your lecture notes before attempting each question.",
        "This assignment doubles as your examination revision.",
        "Draw UML diagrams in correct notation, and check that your models are consistent.",
        "Submit by the date given by your lecturer. Total: 50 marks.",
    ]:
        bullet(doc, it)

    # Section A
    heading(doc, "Section A — Short-Answer Questions  (20 marks)", level=1)
    qa = [
        ("Define object-oriented analysis and object-oriented design.", "3"),
        ("Distinguish an object from a class.", "2"),
        ("Define encapsulation and abstraction, with one example each.", "3"),
        ("Define inheritance and polymorphism.", "2"),
        ("What is UML? State whether it is a modelling or a programming language.", "2"),
        ("Define an actor and a use case.", "2"),
        ("State the meaning of the multiplicity \u201c1 — 0..*\u201d.", "2"),
        ("Differentiate aggregation from composition.", "2"),
        ("Name the four phases of the Unified Process.", "2"),
    ]
    for i, (q, m) in enumerate(qa, 1):
        question(doc, i, q, m)

    # Section B
    heading(doc, "Section B — Application Questions  (20 marks)", level=1)
    para(doc, "Question 1  (12 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "For the mobile money system (users deposit, withdraw, send, receive and check balance; "
              "agents serve customers; administrators manage users and monitor transactions):",
         space_after=4)
    for s in [
        "(a)  Identify the actors.                                               [2]",
        "(b)  Identify the candidate classes and their inheritance.              [3]",
        "(c)  Draw the sequence diagram for \u201cSend Money\u201d.                        [4]",
        "(d)  Draw the state diagram for Transaction.                            [3]",
    ]:
        sub(doc, s)
    para(doc, "", space_after=2)
    para(doc, "Question 2  (8 marks)", size=12, bold=True, color=DARK, space_after=4)
    para(doc, "The sequence diagram calls validateTransaction(), but the class diagram contains "
              "checkTransaction().", space_after=4)
    for s in [
        "(a)  Identify the modelling problem.                                    [2]",
        "(b)  Explain why consistency matters between the diagrams.              [3]",
        "(c)  Describe how the team should resolve the discrepancy.              [3]",
    ]:
        sub(doc, s)

    # Section C
    heading(doc, "Section C — Case Study  (10 marks)", level=1)
    para(doc, "\u201cA university wants an online student registration system. Students can log in, view "
              "courses, check prerequisites, register for courses, drop courses and view their "
              "registration. Administrators can create courses, modify courses, close registration and "
              "view reports.\u201d",
         italic=True, space_after=6)
    for s in [
        "(a)  Identify the actors and the use cases.                             [3]",
        "(b)  Draw the class diagram with relationships and multiplicities.      [4]",
        "(c)  Draw the state diagram for Course.                                 [3]",
    ]:
        sub(doc, s)

    # Marking scheme
    heading(doc, "Marking Scheme", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells
    for j, htxt in enumerate(["Section", "Content", "Marks"]):
        H["set_cell_bg"](head[j], "7B1113")
        H["cell_margins"](head[j])
        p = head[j].paragraphs[0]
        run(p, htxt, size=11, bold=True, color=CREAM, font=H_FONT)
    body = [
        ("A", "Short-answer questions (9 questions)", "20"),
        ("B", "Application questions (12 + 8)", "20"),
        ("C", "Case study — online student registration", "10"),
        ("", "TOTAL", "50"),
    ]
    for i, (a, b, c) in enumerate(body, 1):
        cells = t.rows[i].cells
        for j, txt in enumerate((a, b, c)):
            H["cell_margins"](cells[j])
            if i == 4:
                H["set_cell_bg"](cells[j], "F3E4E4")
            p = cells[j].paragraphs[0]
            bold = (i == 4)
            run(p, txt, size=11, bold=bold, color=(MAROON if bold else INK))
    for j, w in enumerate((Inches(0.9), Inches(3.6), Inches(0.9))):
        t.columns[j].width = w
        for row in t.rows:
            row.cells[j].width = w

    para(doc, "", space_after=4)
    heading(doc, "Submission & Academic Integrity", level=1)
    for it in [
        "Submit a typed document (or neatly handwritten, as directed by your lecturer).",
        "This assignment is individual work. Plagiarism or copying will attract a penalty under university regulations.",
        "Diagrams must use correct UML notation, including multiplicity and visibility.",
        "Cite any sources you consult using the format used in the Week 16 module references.",
    ]:
        bullet(doc, it)

    gold_rule(doc)
    para(doc, "Francis Fweta — ICT Lecturer · Exploits University · BICT 3202 Object-Oriented Analysis and Design",
         size=10, italic=True, color=H["GREY"], align=1)


if __name__ == "__main__":
    print("Built:", build(CFG, content))
